import os
from xml.dom.minidom import Document

import docx
from django.http import HttpResponse, FileResponse
from django.views import View
from django.views.decorators.csrf import csrf_exempt
from django.views.generic.edit import FormView
from django.db.models import Max
from docx.shared import Inches, Cm, Pt
from docx.table import Table

from PythonDjango.settings import MEDIA_DIR
from scripts.import_from_excel import start_import
from worktime.models import Settings
from .forms import PlanForm, UserRegistrationForm
from .models import Plan, Rubric, Direction, Purpose, Plantable
from django.utils.html import escape
from .utils import *

from docxtpl import DocxTemplate


def index(request):
    return render(request, 'plan/index.html')


def ribbon(request):
    s = Settings.objects.filter(field='plantable')[0].value
    table = Plantable.objects.get(pk=int(s))
    rubrics = Rubric.objects.filter(plantable_id=table)
    rubr_id_max = Rubric.objects.aggregate(Max('id'))
    context = {'rubrics': rubrics, 'rubr_id_max': rubr_id_max}
    return render(request, 'plan/ribbon.html', context)


# def updplan (request, id):
#     data = request.POST
#     p = Plan.objects.get(pk=id)
#
#
#
#     context = {}
#     return render(request, 'plan/ribbview.html', context)

def ribbview(request, r_id):
    # form = PlanForm()
    s = Settings.objects.filter(field='plantable')[0].value
    table = Plantable.objects.get(pk=int(s))
    # dir_0 = Direction.objects.filter(plantable_id=table, name='')[0]
    # pur_0 = Purpose.objects.filter(plantable_id=table, name='')[0]
    plans = Plan.objects.filter(plantable_id=table, r_id=r_id)
    rtable = Rubric.objects.get(pk=r_id)

    pl2 = len(Plan.objects.filter(plantable_id=table, r_id=r_id, direction_id__name__gt=''))
    if pl2 == 0:
        dir_no = ' hidden '
    else:
        dir_no = ''

    pl2 = len(Plan.objects.filter(plantable_id=table, r_id=r_id, purpose_id__name__gt=''))
    if pl2 == 0:
        pur_no = ' hidden '
    else:
        pur_no = ''

    # pl2 = pl1.exclude(plantable_id=table).count()
    # if  pl2 == 0:
    #     dir_no = True

    context = {'plans': plans, 'dir_no': dir_no, 'pur_no': pur_no}
    return render(request, 'plan/ribbview.html', context)


def login(request):
    return render(request, 'registration/login.html')


def kostil(s):
    c = s.rfind('/')
    s = s[:c]
    c = s.rfind('/')
    return s[:c + 1] + 'post/', s[:c + 1] + 'postr/'


def make_rubrics(rubrics):
    rubrics_code = ''
    for rubric in rubrics:
        rubrics_code += ' <option id="' + str(rubric.id) + '" level="' + str(
            rubric.riven) + '" >' + rubric.name + '</option >' + ''
    return escape(rubrics_code)


def view(request):
    s = Settings.objects.filter(field='plantable')[0].value
    table = Plantable.objects.get(pk=int(s))
    plans = Plan.objects.filter(plantable_id=table)
    rubrics = Rubric.objects.filter(plantable_id=table)
    rubr_id_max = Rubric.objects.aggregate(Max('id'))
    dir, dir2 = kostil(request.build_absolute_uri())
    context = {'plans': plans, 'rubrics': rubrics, 'dir': dir, 'dir2': dir2, 'rubr_id_max': rubr_id_max}
    # return render(request, reverse_lazy('index'), context)
    return render(request, 'plan/index.html', context)


def post(request, id):
    s = Settings.objects.filter(field='plantable')[0].value
    table = Plantable.objects.get(pk=int(s))
    plans = Plan.objects.filter(plantable_id=table)
    context = {'plans': plans}
    return render(request, 'plan/post.html', context)


def verification_user_permission_for_write(request):
    if request.user.is_authenticated:
        if request.user.has_perm('plan.change_plan'):
            print("Можу міняти")
        if request.user.has_perm('plan.view_plan'):
            print("Можу дивитися")

        # perm = User.user_permissions.

        username = request.user
    # print(username)
    return username


class Post_delete(View):

    def post(self, request, id):
        post = Plan.objects.get(pk=id)
        post.delete()
        context = {}
        return render(request, 'plan/post.html', context)


def postr(request, r_id, num):
    verification_user_permission_for_write(request)
    s = Settings.objects.filter(field='plantable')[0].value
    table = Plantable.objects.get(pk=int(s))

    if (num > 0) and (request.user.has_perm('plan.view_plan')):
        plans = Plan.objects.filter(plantable_id=table, r_id=r_id)
        count = len(plans)
        if count == 0:
            return render(request, 'plan/post_empty.html')
        if num >= count:
            num = count
        plan = plans[num - 1]
        i_id = plan.id
        form = PlanForm(instance=plan)
        if request.method == "POST":
            form.save()
            context = {'num': num, 'count': count, 'form': form, 'r_id': r_id, 'i_id': i_id}
            return render(request, '', context)

        else:
            context = {'num': num, 'count': count, 'form': form, 'r_id': r_id, 'i_id': i_id}
            return render(request, 'plan/post.html', context)
    else:
        if request.user.has_perm('plan.add_plan'):
            plan = Plan()
            plan.r_id = Rubric.objects.get(pk=int(r_id))
            plan.content = ''
            plan.responsible = ''
            plan.termin = ''
            plan.generalization = ''
            plan.note = ''
            plan.direction_id = Direction.objects.filter(name='')[0]
            plan.purpose_id = Purpose.objects.filter(name='')[0]
            max_sort_fild = Plan.objects.filter(r_id=r_id).aggregate(Max('sort'))
            srt = max_sort_fild['sort__max']
            plan.sort = srt + 1

            plan.plantable_id = table
            plan.save()
            form = PlanForm(instance=plan)
            i_id = plan.id
            plans = Plan.objects.filter(r_id=r_id, plantable_id=table)
            count = len(plans)
            num = count
            context = {'num': num, 'count': count, 'form': form, 'r_id': r_id, 'i_id': i_id}
            return render(request, 'plan/post.html', context)


def imp_from_excel(request):
    start_import()
    return render(request, 'plan/import_end.html')


@csrf_exempt
def rib_del_plan(request, id):
    if request.POST and request.is_ajax() and request.user.has_perm('plan.change_plan'):
        print("Вилучамо id=", id)
        Plan.objects.get(pk=id).delete()
        return render(request, 'plan/ribbview.html', {})


@csrf_exempt
def rib_update_plan(request, id, num_field):
    if request.POST and request.is_ajax() and request.user.has_perm('plan.change_plan'):
        data = request.POST
        p = Plan.objects.get(pk=id)
        map = {1: 'id', 2: 'direction_id', 3: 'purpose_id', 4: 'content', 5: 'termin', 6: 'generalization',
               7: 'responsible', 8: 'note'}

        if num_field == 4:
            p.content = data['content']
            p.save()
        elif num_field == 5:
            p.termin = data['termin']
            p.save()
        elif num_field == 6:
            p.generalization = data['generalization']
            p.save()
        elif num_field == 7:
            p.responsible = data['responsible']
            p.save()
        elif num_field == 8:
            p.note = data['note']
            p.save()
        elif num_field == 9:
            s = float(data['sort'].replace(",", "."))
            p.sort = s
            p.save()
        elif num_field == 10:
            p.show = not p.show
            p.save()
        elif num_field == 0:
            p.content = data['content']
            p.termin = data['termin']
            p.generalization = data['generalization']
            p.responsible = data['responsible']
            p.note = data['note']
            p.save()

        # print(map[num_field])

        # if data['direction_id'] == '':
        #     p.direction_id = Direction.objects.get(pk=0)
        # else:
        #     p.direction_id = Direction.objects.get(pk=int(data['direction_id']))
        # if data['purpose_id'] == '':
        #     p.purpose_id = Purpose.objects.get(pk=0)
        # else:
        #     p.purpose_id = Purpose.objects.get(pk=int(data['purpose_id']))
        # p.content = data['content']
        # p.generalization = data['generalization']
        # p.responsible = data['responsible']
        # p.termin = data['termin']
        # p.note = data['note']

    return render(request, 'plan/ribbview.html', {})


@csrf_exempt
def rib_add_plan(request):
    s = Settings.objects.filter(field='plantable')[0].value
    table = Plantable.objects.get(pk=int(s))
    if request.POST and request.is_ajax() and request.user.has_perm('plan.change_plan'):
        data = request.POST
        p = Plan()

        p.content = data['content']
        p.termin = data['termin']
        p.generalization = data['generalization']
        p.responsible = data['responsible']
        p.sort = float(data['sort'].replace(",", "."))
        p.note = data['note']
        p.direction_id = Direction.objects.filter(plantable_id=table, name='')[0]
        p.purpose_id = Purpose.objects.filter(plantable_id=table, name='')[0]

        pltable = data['plantable']
        pt = Plantable.objects.get(pk=int(pltable))
        p.plantable_id = pt

        rtable = data['rtable']
        rt = Rubric.objects.filter(name=rtable)[0]

        p.r_id = rt
        p.save()

        # print(map[num_field])

        # if data['direction_id'] == '':
        #     p.direction_id = Direction.objects.get(pk=0)
        # else:
        #     p.direction_id = Direction.objects.get(pk=int(data['direction_id']))
        # if data['purpose_id'] == '':
        #     p.purpose_id = Purpose.objects.get(pk=0)
        # else:
        #     p.purpose_id = Purpose.objects.get(pk=int(data['purpose_id']))
        # p.content = data['content']
        # p.generalization = data['generalization']
        # p.responsible = data['responsible']
        # p.termin = data['termin']
        # p.note = data['note']

    return render(request, 'plan/ribbview.html', {})


def update_plan(request, id):
    if request.POST and request.is_ajax() and request.user.has_perm('plan.change_plan'):
        data = request.POST
        p = Plan.objects.get(pk=id)
        if data['direction_id'] == '':
            p.direction_id = Direction.objects.get(pk=0)
        else:
            p.direction_id = Direction.objects.get(pk=int(data['direction_id']))
        if data['purpose_id'] == '':
            p.purpose_id = Purpose.objects.get(pk=0)
        else:
            p.purpose_id = Purpose.objects.get(pk=int(data['purpose_id']))
        p.content = data['content']
        p.generalization = data['generalization']
        p.responsible = data['responsible']
        p.termin = data['termin']
        p.note = data['note']
        p.save()
    return render(request, 'plan/post.html', {})


class MyRegisterFormView(FormView):
    form_class = UserRegistrationForm
    success_url = "../../plan/view/"
    # Шаблон, который будет использоваться при отображении представления.
    template_name = "registration/register.html"

    def form_valid(self, form):
        form.save()
        # Функция super( тип [ , объект или тип ] )
        # Возвратите объект прокси, который делегирует вызовы метода родительскому или родственному классу типа .
        return super(MyRegisterFormView, self).form_valid(form)

    def form_invalid(self, form):
        return super(MyRegisterFormView, self).form_invalid(form)


def export_word(request):
    #     http://www.ilnurgi1.ru/docs/python/modules_user/docx.html
    #   https://python-docx.readthedocs.io/en/latest/user/quickstart.html
    # https://python-docx.readthedocs.io/en/latest/user/styles-understanding.html
    # http://qaru.site/questions/1673329/python-docx-how-to-set-cell-width-in-tables
    # http://qaru.site/questions/2267905/modify-docx-page-margins-with-python-docx

    # Границі комірок таблиці
    # https://issue.life/questions/33069697

    document = docx.Document()

    # changing the page margins
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    s = Settings.objects.filter(field='plantable')[0].value
    plantable = Plantable.objects.get(pk=int(s))

    tree = rubric_tree()
    for i, r in enumerate(tree):
        if (tree[i]['riven'] == 1) and (tree[i]['n_r'] != 1):
            document.add_page_break()
        document.add_heading(tree[i]['text'], level=tree[i]['riven'])
        if i == 7:
            break
        print(tree[i]['text'])
        r_id = tree[i]['r_id']
        rubr = Rubric.objects.get(pk=r_id)
        plans = Plan.objects.filter(plantable_id=plantable, r_id=rubr, show=True)
        len_plan = len(plans)
        if len_plan > 0:
            # Будуємо таблицю

            colCount = 8
            rowCount = len_plan + 1
            dir_ = True
            pur_ = True
            pl2 = len(Plan.objects.filter(plantable_id=plantable, r_id=rubr, show=True, direction_id__name__gt=''))
            if pl2 == 0:
                colCount -= 1
                dir_ = False

            pl2 = len(Plan.objects.filter(plantable_id=plantable, r_id=rubr, show=True, purpose_id__name__gt=''))
            if pl2 == 0:
                colCount -= 1
                pur_ = False

            # document.add_paragraph(str(len_plan))
            table = document.add_table(rows=rowCount, cols=colCount)
            table.autofit = False
            # table.rows.

            if (not dir_) and (not pur_):
                c_work = 1
                t_work = 2
                g_work = 3
                r_work = 4
                n_work = 5
            elif (not dir_) and (pur_):
                p_work = 1
                c_work = 2
                t_work = 3
                g_work = 4
                r_work = 5
                n_work = 6
            elif (dir_) and (not pur_):
                d_work = 1
                c_work = 2
                t_work = 3
                g_work = 4
                r_work = 5
                n_work = 6
            else:
                d_work = 1
                p_work = 2
                c_work = 3
                t_work = 4
                g_work = 5
                r_work = 6
                n_work = 7

            cell = table.cell(0, 0)
            # cell.text = '№'
            cell.paragraphs[0].add_run('№').bold = True

            if dir_:
                cell = table.cell(0, d_work)
                cell.paragraphs[0].add_run('Напрямки роботи').bold = True

            if pur_:
                cell = table.cell(0, p_work)
                cell.paragraphs[0].add_run('Мета роботи').bold = True

            cell = table.cell(0, c_work)
            # cell.text = 'Зміст роботи'
            cell.paragraphs[0].add_run('Зміст роботи').bold = True

            cell = table.cell(0, t_work)
            # cell.text = 'Строки виконання'
            cell.paragraphs[0].add_run('Строки виконання').bold = True

            cell = table.cell(0, g_work)
            # cell.text = 'Форма узагальнення'
            cell.paragraphs[0].add_run('Форма узагальнення').bold = True

            cell = table.cell(0, r_work)
            # cell.text = 'Відповідальні'
            cell.paragraphs[0].add_run('Відповідальні').bold = True

            cell = table.cell(0, n_work)
            # cell.text = 'При- мітка'
            cell.paragraphs[0].add_run('При- мітка').bold = True

            # for row in range(1, rowCount):
            row = 0
            for plan in plans:
                print(str(row) + " (" + str(rowCount) + ")")
                row += 1
                cell = table.cell(row, 0)
                cell.text = str(row)

                cell = table.cell(row, n_work)
                cell.text = ''
                if dir_:
                    cell = table.cell(row, d_work)
                    cell.text = plan.direction_id.name
                if pur_:
                    cell = table.cell(row, p_work)
                    cell.text = plan.purpose_id


                cell = table.cell(row, c_work)
                content = plan.content

                # TODO тут не працює
                cell.text = content.replace('\n\n', '\n')

                cell = table.cell(row, t_work)
                cell.text = plan.termin

                cell = table.cell(row, g_work)
                cell.text = plan.generalization

                cell = table.cell(row, r_work)
                cell.text = plan.responsible

                cell = table.cell(row, n_work)

                note = plan.note
                id_aa = str(plan.id)
                p2 = cell.add_paragraph(note)
                p1 = cell.add_paragraph(id_aa)

                p1.runs[0].font.size = Pt(4)

            table.style = 'Table Grid'

            set_col_widths(table, colCount)

            # Зміна розміру швифта в першому рядку

            for col in table.rows[0].cells:
                paragraph = col.paragraphs[0]
                run = paragraph.runs
                font = run[0].font
                font.size = Pt(10)  # font size = 10



    filename = os.path.join(MEDIA_DIR, 'report', 'plan.docx')  # r'd:/MyDoc/PythonDjango/plan.docx'

    document.save(filename)
    # return render(request, 'plan/ribbon.html', {})
    return FileResponse(open(filename, 'rb'), as_attachment=True)


def set_col_widths(table, colCount):
    # https://python-docx.readthedocs.io/en/latest/api/table.html

    if colCount == 6:
        widths = (Inches(0.4), Inches(3.5), Inches(0.9), Inches(1.1), Inches(1.2), Inches(0.6))
        # widths = (10, 200, 20, 20, 20, 20)
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width
    elif colCount == 7:
        widths = (Inches(0.4), Inches(2), Inches(3.5), Inches(0.9), Inches(1.1), Inches(1.2), Inches(0.6))
        # widths = (10, 200, 20, 20, 20, 20)
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width
    elif colCount == 8:
        widths = (Inches(0.4), Inches(2), Inches(2), Inches(3.5), Inches(1.1), Inches(0.9), Inches(1.2), Inches(0.6))
        # widths = (10, 200, 20, 20, 20, 20)
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width
