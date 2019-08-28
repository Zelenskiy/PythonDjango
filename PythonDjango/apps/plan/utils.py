import os

import docx
from django.db.models import Max
from django.shortcuts import get_object_or_404, render, redirect

# class ObjectDetailMixin:
#     model = None
#     template = None
#
#     def get(self, request, slug):
#         obj = get_object_or_404(self.model, slug__iexact=slug)
#         return render(request, self.template, content={self.model.__name__.lower(): obj})
#
# class ObjectCreateMixin:
#     model_form = None
#     template = None
#
#     def gen(self, request):
#         form = self.model_form()
#         return render(request, self.template, content={'form': form})
#
#     def post(self, request):
#         bound_form = self.model_form(request.POST)
#         if bound_form.is_valid():
#             new_obj = bound_form.save()
#             return redirect(new_obj)
#         return render(request, self.template, content={'form': bound_form})


# Експорт плану у Word


# Формування дерева розділів
from docx.shared import Pt, Cm, Inches

from PythonDjango.settings import MEDIA_DIR
from plan.models import Rubric, Plantable, Plan
from worktime.models import Settings


# def tree(r):
#
#
#     return 1

def export_plan_to_word():
    #     http://www.ilnurgi1.ru/docs/python/modules_user/docx.html
    #   https://python-docx.readthedocs.io/en/latest/user/quickstart.html
    # https://python-docx.readthedocs.io/en/latest/user/styles-understanding.html
    # http://qaru.site/questions/1673329/python-docx-how-to-set-cell-width-in-tables
    # http://qaru.site/questions/2267905/modify-docx-page-margins-with-python-docx

    # Границі комірок таблиці
    # https://issue.life/questions/33069697

    document = docx.Document()

    # changing the page margins
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    s = Settings.objects.filter(field='plantable')[0].value
    plantable = Plantable.objects.get(pk=int(s))

    tree = rubric_tree()
    for i, r in enumerate(tree):
        if (tree[i]['riven'] == 1) and (tree[i]['n_r'] != 1):
            document.add_page_break()
        document.add_heading(tree[i]['text'], level=tree[i]['riven'])
        # if i == 7:
        #     break
        print(tree[i]['text'])
        r_id = tree[i]['r_id']
        rubr = Rubric.objects.get(pk=r_id)
        plans = Plan.objects.filter(plantable_id=plantable, r_id=rubr, show=True)
        len_plan = len(plans)
        if len_plan > 0:
            # Будуємо таблицю

            colCount = 8
            rowCount = len_plan + 1
            dir_ = True
            pur_ = True
            pl2 = len(Plan.objects.filter(plantable_id=plantable, r_id=rubr, show=True, direction_id__name__gt=''))
            if pl2 == 0:
                colCount -= 1
                dir_ = False

            pl2 = len(Plan.objects.filter(plantable_id=plantable, r_id=rubr, show=True, purpose_id__name__gt=''))
            if pl2 == 0:
                colCount -= 1
                pur_ = False

            # document.add_paragraph(str(len_plan))
            table = document.add_table(rows=rowCount, cols=colCount)
            table.autofit = False
            # table.rows.

            if (not dir_) and (not pur_):
                c_work = 1
                t_work = 2
                g_work = 3
                r_work = 4
                n_work = 5
            elif (not dir_) and (pur_):
                p_work = 1
                c_work = 2
                t_work = 3
                g_work = 4
                r_work = 5
                n_work = 6
            elif (dir_) and (not pur_):
                d_work = 1
                c_work = 2
                t_work = 3
                g_work = 4
                r_work = 5
                n_work = 6
            else:
                d_work = 1
                p_work = 2
                c_work = 3
                t_work = 4
                g_work = 5
                r_work = 6
                n_work = 7

            cell = table.cell(0, 0)
            # cell.text = '№'
            cell.paragraphs[0].add_run('№').bold = True

            if dir_:
                cell = table.cell(0, d_work)
                cell.paragraphs[0].add_run('Напрямки роботи').bold = True

            if pur_:
                cell = table.cell(0, p_work)
                cell.paragraphs[0].add_run('Мета роботи').bold = True

            cell = table.cell(0, c_work)
            # cell.text = 'Зміст роботи'
            cell.paragraphs[0].add_run('Зміст роботи').bold = True

            cell = table.cell(0, t_work)
            # cell.text = 'Строки виконання'
            cell.paragraphs[0].add_run('Строки виконання').bold = True

            cell = table.cell(0, g_work)
            # cell.text = 'Форма узагальнення'
            cell.paragraphs[0].add_run('Форма узагальнення').bold = True

            cell = table.cell(0, r_work)
            # cell.text = 'Відповідальні'
            cell.paragraphs[0].add_run('Відповідальні').bold = True

            cell = table.cell(0, n_work)
            # cell.text = 'При- мітка'
            cell.paragraphs[0].add_run('При- мітка').bold = True

            # for row in range(1, rowCount):
            row = 0
            for plan in plans:
                print(str(row) + " (" + str(rowCount) + ")")
                row += 1
                cell = table.cell(row, 0)
                cell.text = str(row)

                cell = table.cell(row, n_work)
                cell.text = ''
                if dir_:
                    cell = table.cell(row, d_work)
                    cell.text = plan.direction_id.name
                if pur_:
                    cell = table.cell(row, p_work)
                    cell.text = plan.purpose_id


                cell = table.cell(row, c_work)
                content = plan.content

                # TODO тут не працює
                cell.text = content.replace('\n\n', '\n')

                cell = table.cell(row, t_work)
                cell.text = plan.termin

                cell = table.cell(row, g_work)
                cell.text = plan.generalization

                cell = table.cell(row, r_work)
                cell.text = plan.responsible

                cell = table.cell(row, n_work)

                note = plan.note
                id_aa = str(plan.id)
                p2 = cell.add_paragraph(note)
                p1 = cell.add_paragraph(id_aa)

                p1.runs[0].font.size = Pt(4)

            table.style = 'Table Grid'

            set_col_widths(table, colCount)

            # Зміна розміру швифта в першому рядку

            for col in table.rows[0].cells:
                paragraph = col.paragraphs[0]
                run = paragraph.runs
                font = run[0].font
                font.size = Pt(10)  # font size = 10
    filename = os.path.join(MEDIA_DIR, 'report', 'plan.docx')  # r'd:/MyDoc/PythonDjango/plan.docx'
    document.save(filename)
    return filename


def set_col_widths(table, colCount):
    # https://python-docx.readthedocs.io/en/latest/api/table.html

    if colCount == 6:
        widths = (Inches(0.4), Inches(3.5), Inches(0.9), Inches(1.1), Inches(1.2), Inches(0.6))
        # widths = (10, 200, 20, 20, 20, 20)
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width
    elif colCount == 7:
        widths = (Inches(0.4), Inches(2), Inches(3.5), Inches(0.9), Inches(1.1), Inches(1.2), Inches(0.6))
        # widths = (10, 200, 20, 20, 20, 20)
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width
    elif colCount == 8:
        widths = (Inches(0.4), Inches(2), Inches(2), Inches(3.5), Inches(1.1), Inches(0.9), Inches(1.2), Inches(0.6))
        # widths = (10, 200, 20, 20, 20, 20)
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width

def rubric_tree():
    r_tree = {}
    s = Settings.objects.filter(field='plantable')[0].value
    table = Plantable.objects.get(pk=int(s))

    rubrics = Rubric.objects.filter(plantable_id=table)

    # max_id = Rubric.objects.aggregate(Max('id'))
    i = 0
    for rubric in rubrics:
        if rubric.riven == 1:
            child = {}
            child['text'] = (str(rubric.n_r) + '. ' + rubric.name)
            child['riven'] = 1
            child['n_r'] = rubric.n_r
            child['r_id'] = rubric.id
            r_tree[i] = child
            i += 1
            for rubric2 in rubrics:
                if rubric2.riven == 2 and rubric2.id_owner_id == rubric.id:
                    child = {}
                    child['text'] = (str(rubric.n_r) + '. ' +str(rubric2.n_r) + '. ' + rubric2.name)
                    child['riven'] = 2
                    child['r_id'] = rubric2.id
                    r_tree[i] = child
                    i += 1
                    for rubric3 in rubrics:
                        if rubric3.riven == 3 and rubric3.id_owner_id == rubric2.id:
                            child = {}
                            child['text'] = (str(rubric.n_r) + '. ' + str(rubric2.n_r) + '. ' +str(rubric3.n_r) + '. ' + rubric3.name)
                            child['riven'] = 3
                            child['r_id'] = rubric3.id
                            r_tree[i] = child
                            i += 1
    count_r = i
    return r_tree


"""
{ %
for rubric in rubrics %}
{ % if rubric.riven == 1 %}
< option
id = "{{ rubric.id }}"
owner = "{{ rubric.id_owner_id }}"
hidden_child = true
r = "0" >
{{rubric.n_r}}. & nbsp;
{{rubric.name}}
< / option >
{ %
for rubric2 in rubrics %}
{ % if rubric2.riven == 2 and rubric2.id_owner_id == rubric.id %}
< option
id = "{{ rubric2.id }}"
owner = "{{ rubric2.id_owner_id }}"
hidden_child = true
r = "0" >
& nbsp; & nbsp;
{{rubric.n_r}}. & nbsp;
{{rubric2.n_r}}. & nbsp;
{{rubric2.name}}
< / option >

{ %
for rubric3 in rubrics %}
{ % if rubric3.riven == 3 and rubric3.id_owner_id == rubric2.id %}
< option
id = "{{ rubric3.id }}"
owner = "{{ rubric3.id_owner_id }}"
hidden_child = true
r = "0" >
& nbsp; & nbsp;
& nbsp; & nbsp;
{{rubric.n_r}}. & nbsp;
{{rubric2.n_r}}. & nbsp;
{{rubric3.n_r}}. & nbsp;
{{rubric3.name}}
< / option >
{ % endif %}
{ % endfor %}
{ % endif %}
{ % endfor %}
{ % endif %}
{ % endfor %}
"""
